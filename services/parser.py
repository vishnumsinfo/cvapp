"""Extract clean text from PDF / DOCX before sending to Claude (token saving)."""
import os
import pdfplumber
from docx import Document


class UnsupportedFileError(Exception):
    pass


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _clean(_from_pdf(file_path))
    if ext in (".docx", ".doc"):
        return _clean(_from_docx(file_path))
    raise UnsupportedFileError(f"Unsupported file type: {ext}")


def _from_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _from_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _clean(text: str) -> str:
    # collapse excessive whitespace to conserve tokens
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    return "\n".join(lines)
